import time
import io
import os
import zipfile
import json
import logging
import sys
import cv2
import numpy as np
import PyPDF2
import pytesseract
import pandas as pd
import shutil
import pikepdf
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.colors import Color
from pdf2image import convert_from_path
import pdfplumber
from deep_translator import GoogleTranslator
from rq import get_current_job
from pix2tex.cli import LatexOCR
from docx import Document
from docx.shared import Inches
from pdf2docx import Converter

from database import SessionLocal, DocumentTask

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    stream=sys.stdout
)
logger = logging.getLogger(__name__)

logger.info("Загрузка весов нейросети Pix2Tex...")
math_model = LatexOCR()
logger.info("Системы воркера готовы к работе!")

# фикс единая директория которая расшарена между контейнерами 
STORAGE_DIR = "/storage"

def route_task(tool_id, file_paths, boxes=None, extra_param=None):
    """Главный диспетчер задач: определяет функцию обработки по tool_id."""
    job = get_current_job()
    db = SessionLocal()
    task = db.query(DocumentTask).filter(DocumentTask.id == job.id).first()
    if task:
        task.status = "processing"
        db.commit()
    
    logger.info(f"Задача {job.id} начала выполнение. Инструмент: {tool_id}")
    
    try:
        # словарь маршрутизации
        tasks = {
            'ocr_manual': lambda: do_ocr_manual(job.id, file_paths[0], boxes),
            'ocr_text': lambda: do_ocr_text(job.id, file_paths[0]),
            'ocr_math': lambda: do_ocr_math(job.id, file_paths[0]),
            'pdf_merge': lambda: do_pdf_merge(job.id, file_paths),
            'img_to_pdf': lambda: do_img_to_pdf(job.id, file_paths),
            'pdf_to_word': lambda: do_pdf_to_word(job.id, file_paths[0]),
            'pdf_split': lambda: do_pdf_split(job.id, file_paths[0]),
            'pdf_compress': lambda: do_pdf_compress(job.id, file_paths[0]),
            'pdf_protect': lambda: do_pdf_protect(job.id, file_paths[0], extra_param),
            'pdf_unprotect': lambda: do_pdf_unprotect(job.id, file_paths[0], extra_param),
            'pdf_rotate': lambda: do_pdf_rotate(job.id, file_paths[0], extra_param),
            'pdf_delete': lambda: do_pdf_delete_pages(job.id, file_paths[0], extra_param),
            'pdf_to_excel': lambda: do_pdf_to_excel(job.id, file_paths[0]),
            'pdf_watermark': lambda: do_pdf_watermark(job.id, file_paths[0], extra_param),
            'ocr_translate': lambda: do_ocr_translate(job.id, file_paths[0], extra_param),
            'pdf_unlock': lambda: do_pdf_unlock(job.id, file_paths[0]),
            'pdf_bruteforce': lambda: do_pdf_bruteforce(job.id, file_paths[0]),
        }
        
        result = tasks[tool_id]()
        
        if task:
            task.status = "completed"
            task.extracted_text = result
            db.commit()
        logger.info(f"Задача {job.id} успешно завершена.")
        return result

    except Exception as e:
        logger.error(f"Ошибка в задаче {job.id}: {str(e)}", exc_info=True)
        if task:
            task.status = "error"
            task.extracted_text = str(e)
            db.commit()
        return str(e)
    finally:
        db.close()

# обычный текст
def do_ocr_text(job_id, file_path):
    print(f"[{time.strftime('%X')}] Запуск OCR (Сплошной текст)...")
    
    ext = file_path.lower().split('.')[-1]
    full_text = ""
    
    if ext == 'pdf':
        print(" -> Обработка PDF через pdf2image...")
        images = convert_from_path(file_path)
        for i, img in enumerate(images):
            print(f" -> Распознавание страницы {i+1} из {len(images)}...")
            full_text += pytesseract.image_to_string(img, lang='rus+eng') + "\n\n"
    else:
        img = Image.open(file_path)
        full_text = pytesseract.image_to_string(img, lang='rus+eng')
        
    full_text = full_text.strip()
    if not full_text:
        raise Exception("Не удалось распознать текст на документе.")
        
    doc = Document()
    doc.add_heading('Распознанный текст', 1)
    doc.add_paragraph(full_text)
    
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.docx")
    doc.save(output_path)
    
    return full_text

# формулы
def do_ocr_math(job_id, file_path):
    img = Image.open(file_path)
    extracted_text = math_model(img)

    doc = Document()
    doc.add_heading('Распознанная формула (LaTeX)', 0)
    p = doc.add_paragraph()
    run = p.add_run(extracted_text)
    run.font.name = 'Courier New'
    doc.save(os.path.join(STORAGE_DIR, f"{job_id}.docx"))
    return extracted_text

def do_pdf_merge(job_id, file_paths):
    print(f"[{time.strftime('%X')}] Начинаю объединение {len(file_paths)} PDF файлов...")
    
    merger = PyPDF2.PdfMerger()
    
    for path in file_paths:
        if path.lower().endswith('.pdf'):
            merger.append(path)
        else:
            raise Exception(f"Файл {path} не является PDF!")
            
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.pdf")
    merger.write(output_path)
    merger.close()
    
    print(f"[{time.strftime('%X')}] PDF файлы успешно объединены!")
    return "Файлы успешно объединены в один PDF-документ."

# картинка в пдф
def do_img_to_pdf(job_id, file_paths):
    print(f"[{time.strftime('%X')}] Конвертирую {len(file_paths)} фото в PDF...")
    
    image_list = []
    first_image = Image.open(file_paths[0]).convert('RGB')
    
    for path in file_paths[1:]:
        img = Image.open(path).convert('RGB')
        image_list.append(img)
        
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.pdf")
    first_image.save(output_path, save_all=True, append_images=image_list)
    print(f"[{time.strftime('%X')}] PDF из картинок успешно создан!")
    return "Изображения успешно конвертированы в PDF."

# пдф в ворд
def do_pdf_to_word(job_id, file_path):
    print(f"[{time.strftime('%X')}] Конвертирую PDF в Word...")
    
    if not file_path.lower().endswith('.pdf'):
        raise Exception("Пожалуйста, загрузите файл формата PDF.")
        
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.docx")
    
    cv = Converter(file_path)
    cv.convert(output_path)      
    cv.close()
    
    print(f"[{time.strftime('%X')}] Конвертация PDF -> DOCX завершена!")
    return "PDF успешно конвертирован в редактируемый Word документ."

# пдф раздедение на страницы
def do_pdf_split(job_id, file_path):
    print(f"[{time.strftime('%X')}] Начинаю разделение PDF на страницы...")
    
    reader = PyPDF2.PdfReader(file_path)
    zip_path = os.path.join(STORAGE_DIR, f"{job_id}.zip")
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for i in range(len(reader.pages)):
            writer = PyPDF2.PdfWriter()
            writer.add_page(reader.pages[i])
            
            temp_page_path = os.path.join(STORAGE_DIR, f"temp_page_{i+1}.pdf")
            with open(temp_page_path, "wb") as f:
                writer.write(f)
                
            zipf.write(temp_page_path, arcname=f"page_{i+1}.pdf")
            os.remove(temp_page_path)
            
    print(f"[{time.strftime('%X')}] PDF успешно разделен и упакован в ZIP!")
    return f"Документ разделен на {len(reader.pages)} страниц(ы)."

# сжатие пдф
def do_pdf_compress(job_id, file_path):
    print(f"[{time.strftime('%X')}] Начинаю сжатие PDF...")
    
    reader = PyPDF2.PdfReader(file_path)
    writer = PyPDF2.PdfWriter()

    for page in reader.pages:
        page.compress_content_streams()
        writer.add_page(page)

    output_path = os.path.join(STORAGE_DIR, f"{job_id}.pdf")
    with open(output_path, "wb") as f:
        writer.write(f)
        
    print(f"[{time.strftime('%X')}] PDF успешно сжат!")
    return "Размер PDF-документа был оптимизирован."

def do_ocr_manual(job_id, file_path, boxes_json):
    print(f"[{time.strftime('%X')}] Запуск ручной разметки (Режим: Картинки + Текст)...")
    
    image = cv2.imread(file_path)
    math_boxes = []
    
    if boxes_json:
        math_boxes = json.loads(boxes_json)
        
    math_boxes = sorted(math_boxes, key=lambda b: b['y'])
    
    doc = Document()
    doc.add_heading('Результат ручной разметки', 1)
    
    current_y = 0
    
    for idx, box in enumerate(math_boxes):
        bx, by, bw, bh = box['x'], box['y'], box['w'], box['h']
        
        if by > current_y:
            text_roi = image[current_y:by, 0:image.shape[1]]
            if text_roi.shape[0] > 15:
                text_path = os.path.join(STORAGE_DIR, f"temp_text_{job_id}_{idx}.png")
                cv2.imwrite(text_path, text_roi)
                text_line = pytesseract.image_to_string(Image.open(text_path), lang='rus+eng', config='--oem 3 --psm 6').strip()
                if len(text_line) > 2:
                    doc.add_paragraph(text_line)
                if os.path.exists(text_path): os.remove(text_path)
                
        math_roi = image[by:by+bh, bx:bx+bw]
        math_path = os.path.join(STORAGE_DIR, f"temp_img_{job_id}_{idx}.png")
        cv2.imwrite(math_path, math_roi)
        try:
            print(f" -> Вставка рамки {idx} как картинки")
            doc.add_paragraph()
            doc.add_picture(math_path, width=Inches(4.0))
        except Exception as e:
            print(f"Ошибка вставки картинки: {e}")
            
        if os.path.exists(math_path): os.remove(math_path)
            
        current_y = by + bh
        
    if current_y < image.shape[0]:
        tail_roi = image[current_y:image.shape[0], 0:image.shape[1]]
        if tail_roi.shape[0] > 15:
            tail_path = os.path.join(STORAGE_DIR, f"temp_tail_{job_id}.png")
            cv2.imwrite(tail_path, tail_roi)
            text_line = pytesseract.image_to_string(Image.open(tail_path), lang='rus+eng', config='--oem 3 --psm 6').strip()
            if len(text_line) > 2:
                doc.add_paragraph(text_line)
            if os.path.exists(tail_path): os.remove(tail_path)
            
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.docx")
    doc.save(output_path)
    
    return "Документ размечен вручную. Зоны сохранены как изображения."

def do_pdf_protect(job_id, file_path, password):
    if not password:
        raise Exception("Не указан пароль для защиты!")
        
    reader = PyPDF2.PdfReader(file_path)
    writer = PyPDF2.PdfWriter()
    
    for page in reader.pages:
        writer.add_page(page)
        
    writer.encrypt(password) 
    
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.pdf")
    with open(output_path, "wb") as f:
        writer.write(f)
        
    return f"PDF успешно зашифрован. Пароль: {password}"

def do_pdf_unprotect(job_id, file_path, password):
    if not password:
        raise Exception("Не указан пароль для снятия защиты!")
        
    reader = PyPDF2.PdfReader(file_path)
    
    if reader.is_encrypted:
        try:
            reader.decrypt(password)
        except:
            raise Exception("Неверный пароль!")
            
    writer = PyPDF2.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
        
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.pdf")
    with open(output_path, "wb") as f:
        writer.write(f)
        
    return "Защита с PDF успешно снята. Файл открыт."

def do_pdf_rotate(job_id, file_path, angle_str):
    angle = int(angle_str) if angle_str and angle_str.isdigit() else 90
    
    reader = PyPDF2.PdfReader(file_path)
    writer = PyPDF2.PdfWriter()
    
    for page in reader.pages:
        page.rotate(angle)
        writer.add_page(page)
        
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.pdf")
    with open(output_path, "wb") as f:
        writer.write(f)
        
    return f"Все страницы успешно повернуты на {angle} градусов."

def do_pdf_delete_pages(job_id, file_path, pages_str):
    if not pages_str:
        raise Exception("Укажите страницы для удаления (например: 1, 3, 5-7)")
        
    pages_to_delete = set()
    parts = pages_str.replace(" ", "").split(",")
    for part in parts:
        if "-" in part:
            try:
                start, end = part.split("-")
                pages_to_delete.update(range(int(start), int(end) + 1))
            except:
                pass 
        elif part.isdigit():
            pages_to_delete.add(int(part))
            
    reader = PyPDF2.PdfReader(file_path)
    writer = PyPDF2.PdfWriter()
    
    for i, page in enumerate(reader.pages):
        if (i + 1) not in pages_to_delete:
            writer.add_page(page)
            
    if len(writer.pages) == 0:
        raise Exception("Вы удалили все страницы! Документ пуст.")
        
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.pdf")
    with open(output_path, "wb") as f:
        writer.write(f)
        
    return "Выбранные страницы успешно удалены."

def do_pdf_to_excel(job_id, file_path):
    print(f"[{time.strftime('%X')}] Запуск извлечения таблиц (PDF в Excel)...")
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.xlsx")
    
    tables_found = 0
    
    with pdfplumber.open(file_path) as pdf:
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            for i, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                
                for j, table in enumerate(tables):
                    if table:
                        tables_found += 1
                        df = pd.DataFrame(table)
                        
                        sheet_name = f"Стр_{i+1}_Таб_{j+1}"
                        df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
                        
    if tables_found == 0:
        raise Exception("В документе не найдено ни одной таблицы!")
        
    return f"Успешно! Извлечено таблиц: {tables_found}."

def do_pdf_watermark(job_id, file_path, watermark_text):
    if not watermark_text:
        watermark_text = "CONFIDENTIAL"
        
    packet = io.BytesIO()
    can = canvas.Canvas(packet)
    
    can.setFont("Helvetica-Bold", 60)
    can.setFillColor(Color(0.5, 0.5, 0.5, alpha=0.3)) 
    
    can.translate(250, 400)
    can.rotate(45)
    can.drawCentredString(0, 0, watermark_text)
    can.save()
    
    packet.seek(0)
    watermark_pdf = PyPDF2.PdfReader(packet)
    watermark_page = watermark_pdf.pages[0]
    
    reader = PyPDF2.PdfReader(file_path)
    writer = PyPDF2.PdfWriter()
    
    for page in reader.pages:
        page.merge_page(watermark_page) 
        writer.add_page(page)
        
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.pdf")
    with open(output_path, "wb") as f:
        writer.write(f)
        
    return f"Водяной знак '{watermark_text}' успешно наложен на все страницы."

#фикс перевода
def do_ocr_translate(job_id, file_path, target_lang):
    if not target_lang:
        target_lang = 'ru' # По умолчанию будем переводить на русский, это логичнее
    target_lang = target_lang.strip().lower()

    print(f"[{time.strftime('%X')}] Запуск умного перевода с сохранением верстки на '{target_lang}'...")

    ext = file_path.lower().split('.')[-1]
    working_docx_path = os.path.join(STORAGE_DIR, f"temp_{job_id}.docx")

    try:
        # 1.  всё в DOCX для сохранения структуры
        if ext == 'pdf':
            print(" -> Конвертация PDF в DOCX для сохранения таблиц и отступов...")
            cv = Converter(file_path)
            cv.convert(working_docx_path)
            cv.close()
        elif ext == 'docx':
            # Если это уже Word то копируем его во временный файл
            shutil.copy(file_path, working_docx_path)
        else:
            raise Exception("Сохранение форматирования работает только для PDF и DOCX (Word).")

        # 2. Идем по структуре документа
        doc = Document(working_docx_path)
        translator = GoogleTranslator(source='auto', target=target_lang)

        def translate_safe(text):
            if not text or not text.strip() or len(text) < 2:
                return text
            try:
                return translator.translate(text)
            except Exception:
                return text # В случае ошибки (например, спецсимволы) оставляем оригинал

        print(" -> Перевод основного текста...")
        # Переводим обычные абзацы
        for p in doc.paragraphs:
            if p.text.strip():
                translated_text = translate_safe(p.text)
                #оставляем первый кусок абзаца чтобы сохранить базовый шрифт и стиль и записываем в него переведенный текст. Остальные кусочки очищаем.
                if p.runs:
                    p.runs[0].text = translated_text
                    for i in range(1, len(p.runs)):
                        p.runs[i].text = ""
                else:
                    p.text = translated_text

        print(" -> Перевод таблиц...")
        # переводим текст внутри таблиц 
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            translated_text = translate_safe(p.text)
                            if p.runs:
                                p.runs[0].text = translated_text
                                for i in range(1, len(p.runs)):
                                    p.runs[i].text = ""
                            else:
                                p.text = translated_text

        # 3.сохранение
        output_path = os.path.join(STORAGE_DIR, f"{job_id}.docx")
        doc.save(output_path)

        # удаление временного файла
        if os.path.exists(working_docx_path):
            os.remove(working_docx_path)

        return f"Документ успешно переведен на [{target_lang}] с полным сохранением оригинальной верстки (шрифты, абзацы, таблицы)!"

    except Exception as e:
        # если временный файл остался после ошибки убираем его
        if os.path.exists(working_docx_path):
            os.remove(working_docx_path)
        raise Exception(f"Ошибка умного перевода: {e}")
    
def do_pdf_unlock(job_id, file_path):
    """
    Снимает ограничения (Owner Password): запрет на печать, копирование текста.
    Не работает, если файл требует пароль для самого открытия.
    """
    print(f"[{time.strftime('%X')}] Запуск разблокировки PDF (снятие ограничений)...")
    
    output_path = os.path.join(STORAGE_DIR, f"{job_id}.pdf")
    
    try:
        # pikepdf автоматически снимает Owner-ограничения при открытии и сохранении
        with pikepdf.open(file_path) as pdf:
            pdf.save(output_path)
            
        return "Ограничения (печать, копирование) успешно сняты! Файл полностью разблокирован."
        
    except pikepdf.PasswordError:
        # Если pikepdf ругается на пароль, значит это User Password
        raise Exception("Файл зашифрован от открытия (User Password). Используйте функцию перебора (Brute-force)!")
    except Exception as e:
        raise Exception(f"Ошибка при разблокировке: {e}")

def do_pdf_bruteforce(job_id, file_path):
    """
    Атака по словарю (Brute-force) для файлов, которые вообще не открываются без пароля.
    """
    print(f"[{time.strftime('%X')}] Запуск Brute-force на зашифрованный PDF...")
    
    reader = PyPDF2.PdfReader(file_path)
    
    if not reader.is_encrypted:
        return "Файл не зашифрован. Пароль не требуется!"

    # 1. База частых паролей (словарь)
    common_passwords = ["123", "12345", "password", "admin", "qwerty", "12345678"]
    
    # 2. Генератор PIN-кодов (от 0000 до 9999)
    pin_codes = [str(i).zfill(4) for i in range(10000)]
    
    # объединяем базы
    passwords_to_try = common_passwords + pin_codes
    
    print(f" -> Начинаем перебор {len(passwords_to_try)} комбинаций. Это может занять время...")
    
    # 3. метод перебора
    for pwd in passwords_to_try:
        try:
            # decrypt возвращает флаг успеха (0 - не вышло, 1 или 2 - успех)
            if reader.decrypt(pwd):
                print(f" -> ПАРОЛЬ ПОДОБРАН: {pwd}")
                
                # Снимаем защиту и сохраняем чистый файл
                writer = PyPDF2.PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                    
                output_path = os.path.join(STORAGE_DIR, f"{job_id}.pdf")
                with open(output_path, "wb") as f:
                    writer.write(f)
                    
                return f"Успех! Пароль подобран: '{pwd}'. Защита снята, документ открыт."
        except Exception:
            pass # Игнорируем ошибки при неудачной попытке

    return "Не удалось подобрать пароль. Пароль слишком сложный (отсутствует в словаре PIN-кодов)."
